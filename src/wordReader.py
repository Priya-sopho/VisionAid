#pip install python-docx
#pip install pypiwin32 or pywin32
#pip install num2words
# importing required modules
from docx import Document
import speak
import num2words as n2w
import os 
import re
import threading,time
from pynput import keyboard 
import msvcrt

"""
 wordReader to listen the content of word file
 Following are the function alongwith keyboard keys to be pressed to execute:
 1. Pause/Resume - Space
 2. Change Speed - up key to increase speed, down key to decrease speed
 3. Find a keyword in file - f
 4. Jump on a specific para - j
 5. Rewind (Move 10 lines back) - left key
 6. Repeat - r
 7. Skip current para - right key
 8. Quit - Esc key or q 
"""
class wordReader:
	task = ['Start','Pause','Resume','Exit','Rewind','Repeat','Skip current para','Jump on a para',
		'Search','Change speed']

	def __init__(self, file):
		 
		try:
			# creating a file object
			self.FileObj = open(file, 'rb') 
				 
			# creating a document object
			self.document = Document(self.FileObj)
		except:
			speak.say('Unable to open the file')
			os._exit(1)

		#line no.
		self._line = 0
		
		#para no.
		self._para = 0
	
		self.paragraphs = self.document.paragraphs
	
		self.paragraphs = [p for p in self.paragraphs if len(p.text) != 0]		 
		
		#Number of paras in doc file
		self.total_paras = len(self.paragraphs)
		
		# creating a para object
		self.paraObj = self.paragraphs[self._para]
		
		# extracting text from para
		self.lines = self.paraObj.text.splitlines()
		
		#Lock
		self.lock = threading.Lock()  #Initially pause is on
		
		#Pause button
		self.pause_key = 0

		kb = threading.Thread(target=self.listenKeyboard)
		kb.start()
		

	"""
	 Destructor
	"""
	def __del__(self):
	  	self.FileObj.close()

	"""
	 Keyboard listening Callback function
	"""
	def on_press(self,key):
		#Clear the input flush
		while msvcrt.kbhit():
			msvcrt.getch()
			
		try:
			print('You pressed {0}'.format(key))
			if key.char == 'q':
				self.exit()
			if key.char == 'j':
			    self.jump()
			elif key.char == 'f':
			    self.search()
			elif key.char == 'r':
			    self.repeat() 
		        
		except AttributeError:
			if (key == keyboard.Key.space) and self.pause_key:
				self.pause_key = 0
				self.lock.release()
				    #self.resume()
			elif key == keyboard.Key.space:
				self.lock.acquire()
				self.pause()
			elif(key == keyboard.Key.up):
			    speak.change_speed(1)
			elif(key == keyboard.Key.down):
			    speak.change_speed(-1)
			elif(key == keyboard.Key.right):
			    self.skip()
			elif(key == keyboard.Key.left): 
			    self.rewind()
			elif(key == keyboard.Key.esc):
			    self.exit()

	"""
	 Listen to keyboard keys
	"""
	def listenKeyboard(self):
	    # Collect events until released
	    with keyboard.Listener(on_press=self.on_press) as listener:
	        listener.join()


	def  findText(self,s):
		paras = []
		s = re.compile(s,re.I)
		i = 1
		for para in self.paragraphs:
			content = para.text
			#print(content)
			if s.search(content) is not None:
				if i not in paras:
					paras.append(i)
			i += 1		
		return paras


	def menu(self):
		sense.Speak('What do you want to do?')
		for i in range(len(self.task)):
			sense.Speak(n2w.num2words(i+1,to='ordinal')+' '+self.task[i])
		response = int(raw_input())
		if response < len(self.task):
			sense.Speak(self.task[response-1])
		self.action(response)


	"""
	 Start the word reading task.
	"""
	def start(self):
		try:
			
			if self.total_paras>1:
				para = ' paragraphs.'
			else:
				para = ' paragraph.'

			self.lock.acquire()
			speak.say('Welcome to word listening task')
			speak.say('There are '+ n2w.num2words(self.total_paras,to='cardinal') + para)
			self.lock.release()
			self.resume()
		except :
			speak.say('Some Error')
			return
	
	
	"""
	  Pause word reading task
	"""
	def pause(self):
		self.pause_key = 1
		speak.say('Pausing')


	"""
	 Resume word task
	"""
	def resume(self):
		#Wait until pause is not false
		try:
			if(self._para >= self.total_paras):
				speak.say('We reached the end of the doc file')
				self.exit()
			
			#if on first line of para,then tell para number
			if self._line == 0:
				self.lock.acquire()
				speak.say(n2w.num2words(self._para+1, to='ordinal') + ' para') 
				self.lock.release()
			
			# creating a para object
			self.paraObj = self.paragraphs[self._para]
			
			# extracting text from para
			self.lines = self.paraObj.text.splitlines()
			
			if len(self.lines) == 0:
				speak.say('Unable to detect the content')
			
			while self._line < len(self.lines):
				self.lock.acquire()
				speak.say(self.lines[self._line].encode('ascii','ignore').decode('ascii'))
				self._line += 1		
				self.lock.release()
			self._para += 1
			self._line = 0
			self.resume()

		except:
			speak.say("Some error")
			return
	
	

	"""
	 Exit word Reading Task
	"""
	def exit(self):		
		speak.say('Exiting word reading task')
		#Clear the input flush
		while msvcrt.kbhit():
			msvcrt.getch()
		os._exit(1)
		
	

	"""
	  Repeat the last line
	"""	
	def repeat(self):
		self.lock.acquire()
		self._line -= 1
		if(self._line<0):
			self._line = 0
			self._para -= 1
			if(self._para < 0):
				self._para = 0
			#if on first line of para,then tell para number
			if self._line == 0:
				speak.say(n2w.num2words(self._para+1, to='ordinal') + 'para') 
			
			# creating a para object
			self.paraObj = self.paragraphs[self._para]
			
			# extracting text from para
			self.lines = self.paraObj.text.splitlines()
			if len(self.lines) == 0:
				speak.say('Unable to detect the content')
		self.lock.release()
	
	"""
	 Move 10 lines back
	"""
	def rewind(self):
		self.lock.acquire()
		self._line -= 10
		if(self._line<0):
			self._line = 0
			self._para -= 1
			if(self._para < 0):
				self._para = 0
			# creating a para object
			self.paraObj = self.paragraphs[self._para]
			
			# extracting text from para
			self.lines = self.paraObj.text.splitlines()
			
			self._line = len(self.lines) - 10
			if(self._line < 0):
				self._line = 0
			#if on first line of para,then tell para number
			if self._line == 0:
				speak.say(n2w.num2words(self._para+1, to='ordinal') + 'para') 
			if len(self.lines) == 0:
				speak.say('Unable to detect the content')
		self.lock.release()
	

	"""
	 Skip current and move to next para
	"""
	def skip(self):
		self.lock.acquire()
		speak.say('Moving to next para')
		self._para += 1
		self._line = 0
		if(self._para >= self.total_paras):
			speak.say('We reached the end of file')
			self.exit()

		#if on first line of para,then tell para number
		if self._line == 0:
			speak.say(n2w.num2words(self._para+1, to='ordinal') + 'para') 
		
		# creating a para object
		self.paraObj = self.paragraphs[self._para]
			
		# extracting text from para
		self.lines = self.paraObj.text.splitlines()
		if len(self.lines) == 0:
			speak.say('Unable to detect the content')
		self.lock.release()
	
		
	"""
	 Jump on a specific para
	"""		
	def jump(self):
		#Clear the input flush
		while msvcrt.kbhit():
			msvcrt.getch()
		self.lock.acquire()
		speak.say('Tell me the para number to jump on?')
		self._para = int(raw_input())-1
		if(self._para < 0):
			self._para = 0
		if(self._para > self.total_paras):
			self._para = self.total_paras -1
		if(self._para >= self.total_paras):
			speak.say('We reached the end of file')
			self.exit()
		self._line = 0
		#if on first line of para,then tell para number
		if self._line == 0:
			speak.say(n2w.num2words(self._para+1, to='ordinal') + 'para') 
		
		# creating a para object
		self.paraObj = self.paragraphs[self._para]
			
		# extracting text from para
		self.lines = self.paraObj.text.splitlines()
		if len(self.lines) == 0:
			speak.say('Unable to detect the content')
		self.lock.release()
	
	"""
	 Search for a specific word in word file
	"""
	def search(self):
		#Clear the input flush
		while msvcrt.kbhit():
			msvcrt.getch()
		self.lock.acquire()
		speak.say('What you want to search?')
		r = raw_input()
		paras = self.findText(r)
		if len(paras) == 0:
			speak.say(r+' not found!!')
		else :	
			if len(paras)>1:
				para = 'paras.'
			else:
				para = 'para.'
			speak.say(r+ ' found in '+ n2w.num2words(len(paras), to='cardinal') + para)
			for p in paras:
				speak.say('Found at para number'+ n2w.num2words(p+1,to='cardinal'))
		self.lock.release()

	"""
	 Change the rate of speaking up key increase the speed, down key decrease the speed
	"""
	def change_speed(self,r):
		voice.Rate += r
	
#Use with speech commands	
def Task(phrase):
	phrase.lower()
	if phrase == 'start':
		start()
	elif phrase == 'pause':
		pause()
	elif phrase == 'resume':
		resume()
	elif phrase == 'exit':
		exit()
	elif phrase == 'rewind':
		rewind()
	elif phrase == 'repeat':
		repeat()
	elif phrase == 'skip':
		skip()
	elif phrase == 'jump':
		jump()
	elif phrase == 'search':
		search()
	elif phrase == 'change speed':
		change_speed()
	else:
		speak.say('Sorry! Unable to recognize your action')
	
	
speak.say("WORD READER. Give the name of the file: ")
file = raw_input()
if '.docx' not in file:
	file = file + '.docx'
file = os.path.join('word',file)
reader = wordReader(file)
reader.start()		

		
