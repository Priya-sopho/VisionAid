from __future__ import print_function
from docx import Document
import re
from docx.shared import Pt,Inches
from docx.enum.style import WD_STYLE_TYPE
import threading
import os
import speak
from pynput import keyboard

class docEditor:
	def __init__(self):

		speak.say("Word Document. Give the name of the file: ")
		filename = raw_input()
		if '.docx' not in filename:
			filename = filename + '.docx'
		self.filename = filename
		try:
			self.doc = Document(filename)
		except:
		 	try:
		 		#If new file
		 		self.filename = os.path.join("word",self.filename)
		 		self.doc = Document()
		 		self.saveAsDoc()
		 	except:
		 		speak.say('The document cannot be created due to some error.Try to use any other document name!')
		 		self.exit()
		
		self.para = self.doc.add_paragraph('')

		# kb = threading.Thread(target=self.listenKeyboard)
		# kb.start()
		self.menu()
		
	def __del__(self):
		self.doc.close()
		
	#Exit 
	def exit(self):
		speak.say('Exiting Word doc handling Task')
		os._exit(1)

# to get the whole document text for read operation
	def getFullText(self):
		speak.say('Content of file')
		fulltext = []
		for paras in self.doc.paragraphs:
			fulltext.append(paras.text)
		return '\n'.join(fulltext)

# to save the doc as a word document file
	def saveAsDoc(self):
		self.doc.save(self.filename)

# Enter the text to be added
	def enterText(self):
		speak.say("Enter the text: ")
		lines = []
		while True:
			line = raw_input()
			if line:
				lines.append(line)
			else:
				break
		text = '\n'.join(lines)
		return text

# adding new paragraph at the end of the document 
	def addNewPara(self):
		text = self.enterText()
		self.para = self.doc.add_paragraph(text)
		self.saveAsDoc()
		
# add formatted text in the running paragraph
	def formattedtext(self,formats):
		types = formats.split(' ')
		speak.say("Do u want to enter formatted text on a new line? Yes or No : ")
		
		lineBreak = raw_input().upper()
		run = self.para.add_run()
		text = self.enterText()
		run.add_text(text) 
		# style = self.para.style
		# font  = style.font
		if lineBreak == 'Y' or lineBreak == 'YES':
			run.add_break()
		if 'BOLD' or 'B' in types:
			# font.bold = not font.bold
			run.bold = True
		if 'ITALIC' or 'I' in types:
			# font.italic = not font.italic
			run.italic = True
		if 'UNDERLINE' or 'U' in types:
			# font.underline = not font.underline
			run.underline = True
		
		self.saveAsDoc()

# to make a new table 
	def makeTable(self):
		speak.say("Enter the following" + '\n' + 'Rows: ')
		
		row_count = int(raw_input())
		speak.say("Columns: ")
		col_count = int(raw_input())
		table = self.doc.add_table(row_count,col_count)
		table.style = 'Table Grid'
		speak.say("Enter the values in the table row by row: ")
		for row in table.rows:
			for cell in row.cells:
				cell.text = raw_input()	
			print('\n')	
		self.saveAsDoc()

# to change the document's font name and font size
	def changeDocFont(self,fontName, Size):
		style = self.doc.styles['Normal']
		font = style.font
		font.name = fontName
		font.size = Pt(Size)

# to add a heading to the document
	def addHeading(self):
		speak.say("Enter the heading level: ")
		l = int(raw_input())
		text = self.enterText()
		self.doc.add_heading(text,level = l)
		self.saveAsDoc()

# to add a picture to the doc file
	def addPicture(self):
		speak.say("Enter the name of the Picture (with extension i.e. .png or .jpeg): ")
		pic_name = raw_input()
		pic_name = os.path.join("image",pic_name)
		speak.say("Enter the picture width in inches: ")
		pic_width = float(raw_input())
		pic = self.doc.add_picture(pic_name, width = Inches(pic_width))
		self.saveAsDoc()

#to find the paragraph with the particular word
	def search_first_para(self):
		speak.say("Enter the word to search in paragraphs: ")
		word = raw_input()
		word = re.compile(word,re.I)
		for paras in self.doc.paragraphs:
			if word.search(paras.text) is not None:
				return paras
		return None

#to delete a particular paragraph from the document
	def del_para(self,paragraph):
		p = paragraph._element
		if len(p):
			p.getparent().remove(p)
			p._p = p._element = None
			self.para = self.doc.paragraphs[-1]
			self.saveAsDoc()
			return True
		return False

	def menu(self):
		Flag = True
		while Flag == True:
			speak.say('Do you want to listen to menu? (Y or N)')
			choice = raw_input().upper()
			if(choice == 'Y' or choice == 'YES'):
				speak.say('1. Read the doc' + '\n' +
						'2. Add a Heading' + '\n' +
						'3. Add new Para' + '\n' +
						'4. Add formatted text' +'\n' +
						'5. Change Font style of Doc' + '\n'
						'6. Add Table' + '\n'
						'7. Add Picture ' + '\n' 
						'8. Delete a paragraph containing a particular word ' + '\n' 
						'9. Delete the last paragraph ')
			speak.say("Enter your choice number: ")
			flag2 = False
			try:
				ch = int(raw_input())
				flag2 = True
			except:
				speak.say('Oops! you have not entered a number.')

			if flag2:
				if ch == 1:
					try:
						speak.say(self.getFullText())    # to display the text
					except:
						speak.say("Error in reading the document!")
				elif ch == 2:
					try:
						self.addHeading()
					 	speak.say("Heading added!")
					except:
						speak.say("Error in adding the heading!")
				elif ch == 3:
					try:
						self.addNewPara()
						speak.say("Paragraph added!")
					except:
						speak.say("Error in adding new Paragraph!")
				elif ch == 4:
					while flag2 == True:
						try:
							new_format = ''
							speak.say('Enter your choice: bold(B), italic(I) or underline(U) (seperated by spaces): ')
							new_format = raw_input().upper()
							self.formattedtext(new_format)
							speak.say("Do You want to add more formatted text? y or n: ")
							Continue = raw_input().upper()
							if Continue == 'N' or Continue == 'NO':
								flag2 = False
						except:
							speak.say("Error in adding formatted text!")
				elif ch == 5:
					speak.say("Enter the font name and font size (space seperated):")
					name, size = raw_input().split(' ')
					size = int(size)
					self.changeDocFont(name, size)
				elif ch == 6:
					try:
						self.makeTable()
						speak.say("Table added!")
					except:
						speak.say("Error in adding table!")
				elif ch == 7:
					try:
						self.addPicture()
						speak.say("Picture added!")
					except:
						speak.say("Error in adding picture!")
				elif ch == 8:
					try:
						para = self.search_first_para()
						if para is not None:
							r = self.del_para(para)
							if r:
								speak.say("Paragraph deleted!")
							else:
								speak.say("Unable to delete the paragraph")
						else:
							speak.say("Paragraph with searched word doesn't exist")
					except:
						speak.say("Error in deleting paragraph!")
				elif ch == 9:
					try:
						r = self.del_para(self.para)
						if r:
							speak.say("Last Paragraph deleted!")
						else:
							speak.say("No last Paragraph exist")
					except:
						speak.say("Error in deleting paragraph!")
				else:
					speak.say('Invalid choice')
			speak.say("Do You want to perform another operation? y or n: ")		
			another_operation = raw_input().upper()
			if another_operation == 'N' or another_operation == 'No':
				Flag = False
				self.exit()

	# """
	#  Keyboard listening callback function
	# """
	# def on_press(self,key):
	# 	try:
	# 		speak.say('{0}'.format(key.char))
	# 	except AttributeError:
	# 		speak.say('{0}'.format(key))
				
	# """
	#  Listen to keyboard keys
	# """
	# def listenKeyboard(self):
	#     # Collect events until released
	#     with keyboard.Listener(on_press=self.on_press) as listener:
	#         listener.join()
	     


document = docEditor()