from __future__ import print_function
from docx import Document
import re
from docx.shared import Pt,Inches
from docx.enum.style import WD_STYLE_TYPE
import threading
import os
import win32com.client as wincl
from pynput import keyboard
voice = wincl.Dispatch("SAPI.SpVoice")

class docEditor:
	def __init__(self):

		voice.Speak("Word Document. Give the name of the file: ")
		filename = raw_input()
		if '.docx' not in filename:
			filename = filename + '.docx'
		self.filename = filename
		try:
			self.doc = Document(filename)
		except:
			voice.Speak('The document cannot be created due to some error.Try to use any other document name!')
			self.exit()
		
		self.para = self.doc.add_paragraph('')

		kb = threading.Thread(target=self.listenKeyboard)
		kb.start()

		self.menu()
		
	def __del__(self):
		self.doc.close()

	
	#Exit 
	def exit(self):
		os._exit(1)

#clear the input flush
	def clear_input(self):
		while msvcrt.kbhit():
			msvcrt.getch()

# to get the whole document text for read operation
	def getFullText(self):
		fulltext = []
		for paras in self.doc.paragraphs:
			fulltext.append(paras.text)
		return '\n'.join(fulltext)

# to save the doc as a word document file
	def saveAsDoc(self):
		self.doc.save(self.filename)

# Enter the text to be added
	def enterText(self):
		voice.Speak("Enter the text: ")
		self.clear_input()
		lines = []
		while True:
			line = raw_input()
			if line:
				lines.append(line)
			else:
				break
		text = '\n'.join(lines)
		return text

# adding new paragraph at the end of the document 
	def addNewPara(self):
		text = self.enterText()
		self.para = self.doc.add_paragraph(text)
		self.saveAsDoc()
		
# add formatted text in the running paragraph
	def formattedtext(self,formats):
		types = formats.split(' ')
		voice.Speak("Do u want to enter formatted text on a new line? Yes or No : ")
		self.clear_input()
		lineBreak = raw_input().upper()
		run = self.para.add_run()
		text = self.enterText()
		run.add_text(text) 
		# style = self.para.style
		# font  = style.font
		if lineBreak == 'Y' or lineBreak == 'YES':
			run.add_break()
		if 'BOLD' or 'B' in types:
			# font.bold = not font.bold
			run.bold = True
		if 'ITALIC' or 'I' in types:
			# font.italic = not font.italic
			run.italic = True
		if 'UNDERLINE' or 'U' in types:
			# font.underline = not font.underline
			run.underline = True
		
		self.saveAsDoc()

# to make a new table 
	def makeTable(self):
		voice.Speak("Enter the following" + '\n' + 'Rows: ')
		self.clear_input()
		row_count = int(raw_input())
		voice.Speak("Columns: ")
		self.clear_input()
		col_count = int(raw_input())
		table = self.doc.add_table(row_count,col_count)
		table.style.name = 'TableGrid'
		voice.Speak("Enter the values in the table row by row: ")
		self.clear_input()
		for row in table.rows:
			for cell in row.cells:
				cell.text = raw_input()	
			print('\n')	
		print('\n')
		for row in table.rows:
			for cell in row.cells:
				print(cell.text , end = ' ')
			print('\n')
		print('\n')
		self.saveAsDoc()

# to change the document's font name and font size
	def changeDocFont(self,fontName, Size):
		style = self.doc.styles['Normal']
		font = style.font
		font.name = fontName
		font.size = Pt(Size)

# to add a heading to the document
	def addHeading(self):
		voice.Speak("Enter the heading level: ")
		self.clear_input()
		l = int(raw_input())
		text = self.enterText()
		self.doc.add_heading(text, level = l)
		self.saveAsDoc()

# to add a picture to the doc file
	def addPicture(self):
		voice.Speak("Enter the name of the Picture (with extension i.e. .png or .jpeg): ")
		self.clear_input()
		pic_name = raw_input()
		voice.Speak("Enter the picture width in inches: ")
		self.clear_input()
		pic_width = float(raw_input())
		pic = self.doc.add_picture(pic_name, width = Inches(pic_width))
		self.saveAsDoc()

#to find the paragraph with the particular word
	def search_first_para(self):
		voice.Speak("Enter the word: ")
		self.clear_input()
		word = raw_input()
		word = re.compile(word,re.I)
		for paras in self.doc.paragraphs:
			if word.search(paras.text) is not None:
				return paras

#to delete a particular paragraph from the document
	def del_para(self,paragraph):
		p = paragraph._element
		p.getparent().remove(p)
		p._p = p._element = None
		self.para = self.doc.paragraphs[-1]
		self.saveAsDoc()

	def menu(self):
		Flag = True
		while Flag == True:
			voice.Speak('Do you want to listen to menu? (Y or N)')
			self.clear_input()
			choice = raw_input().upper()
			if(choice == 'Y' or choice == 'YES'):
				voice.Speak('1. To read the document' + '\n' +
						'2. To add a Heading' + '\n' +
						'3. To write a new paragraph on the document' + '\n' +
						'4. To add a new formatted text' +'\n' +
						'5. To change the font style and font name of the document' + '\n'
						'6. To add a new Table and then display it' + '\n'
						'7. To add Picture to the document ' + '\n' 
						'8. To delete a paragraph containing a particular word ' + '\n' 
						'9. To delete the last paragraph ')
			voice.Speak("Enter your choice number: ")
			try:
				self.clear_input()
				ch = int(raw_input())
				flag2 = True
			except:
				voice.Speak('Oops! you have not entered a number.')

			if flag2:
				if ch == 1:
					try:
						voice.Speak(self.getFullText())    # to display the text
					except:
						voice.Speak("Error in reading the document!")
				elif ch == 2:
					try:
						self.addHeading()
						voice.Speak("Heading added!")
					except:
						voice.Speak("Error in adding the heading!")
				elif ch == 3:
					try:
						self.addNewPara()
						voice.Speak("Paragraph added!")
					except:
						voice.Speak("Error in adding new Paragraph!")
				elif ch == 4:
					while flag2 == True:
						try:
							new_format = ''
							voice.Speak('Enter your choice: bold(B), italic(I) or underline(U) (seperated by spaces): ')
							self.clear_input()
							new_format = raw_input().upper()
							self.formattedtext(new_format)
							voice.Speak("Do You want to add more formatted text? y or n: ")
							self.clear_input()
							Continue = raw_input().upper()
							if Continue == 'N' or Continue == 'NO':
								flag2 = False
						except:
							voice.Speak("Error in adding formatted text!")
				elif ch == 5:
					voice.Speak("Enter the font name and font size (space seperated):")
					self.clear_input()
					name, size = raw_input().split(' ')
					size = int(size)
					self.changeDocFont(name, size)
				elif ch == 6:
					try:
						self.makeTable()
						voice.Speak("Table added!")
					except:
						voice.Speak("Error in adding table!")
				elif ch == 7:
					try:
						self.addPicture()
						voice.Speak("Picture added!")
					except:
						voice.Speak("Error in adding picture!")
				elif ch == 8:
					try:
						self.del_para(self.search_first_para())
						voice.Speak("Paragraph deleted!")
					except:
						voice.Speak("Error in deleting paragraph!")
				elif ch == 9:
					try:
						self.del_para(self.para)
						voice.Speak("Last Paragraph deleted!")
					except:
						voice.Speak("Error in deleting paragraph!")
				else:
					voice.Speak('Invalid choice')
			voice.Speak("Do You want to perform another operation? y or n: ")		
			another_operation = raw_input().upper()
			if another_operation == 'N' or another_operation == 'No':
				Flag = False
				self.exit()

	"""
	 Keyboard listening callback function
	"""
	def on_press(self,key):
		try:
			voice.Speak('{0}'.format(key.char))
		except AttributeError:
			voice.Speak('{0}'.format(key))
				
	"""
	 Listen to keyboard keys
	"""
	def listenKeyboard(self):
	    # Collect events until released
	    with keyboard.Listener(on_press=self.on_press) as listener:
	        listener.join()
	     


document = docEditor()